"""Report Download Utilities for Multiple Formats"""

import io
import streamlit as st
from typing import Dict, Any, Optional
from datetime import datetime
import os
import tempfile

# Import required libraries
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    st.error("python-docx 라이브러리가 필요합니다. pip install python-docx")

try:
    from fpdf import FPDF
except ImportError:
    st.error("fpdf2 라이브러리가 필요합니다. pip install fpdf2")

try:
    import hwp
except ImportError:
    hwp = None

class ReportDownloader:
    """Report download utility for multiple formats"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx']
        if hwp:
            self.supported_formats.append('hwp')
    
    def create_word_report(self, content: str, title: str = "서울 상권 분석 보고서") -> io.BytesIO:
        """Create a Word document from the report content."""
        try:
            document = Document()
            
            # Add title
            title_paragraph = document.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            document.add_paragraph(f"생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}")
            document.add_paragraph("")
            
            # Parse content
            self._parse_content_to_word(document, content)
            
            # Save to BytesIO
            bio = io.BytesIO()
            document.save(bio)
            bio.seek(0)
            return bio
            
        except Exception as e:
            st.error(f"Word 문서 생성 오류: {e}")
            return None
    
    def create_pdf_report(self, content: str, title: str = "서울 상권 분석 보고서") -> io.BytesIO:
        """Create a PDF document from the report content."""
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Set Korean font
            try:
                # Try to use Korean font
                font_path = self._get_korean_font_path()
                if font_path and os.path.exists(font_path):
                    pdf.add_font('Korean', '', font_path, uni=True)
                    pdf.set_font('Korean', '', 12)
                else:
                    pdf.set_font('Arial', '', 12)
                    st.warning("한글 폰트를 찾을 수 없어 PDF에서 한글이 깨질 수 있습니다.")
            except:
                pdf.set_font('Arial', '', 12)
            
            # Add title
            pdf.set_font('Korean' if 'Korean' in pdf.fonts else 'Arial', 'B', 16)
            pdf.cell(0, 15, title, ln=True, align='C')
            pdf.ln(5)
            
            # Add metadata
            pdf.set_font('Korean' if 'Korean' in pdf.fonts else 'Arial', '', 10)
            pdf.cell(0, 8, f"생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}", ln=True)
            pdf.ln(5)
            
            # Parse content
            self._parse_content_to_pdf(pdf, content)
            
            # Generate PDF
            pdf_output = pdf.output(dest='S').encode('latin-1')
            bio = io.BytesIO(pdf_output)
            bio.seek(0)
            return bio
            
        except Exception as e:
            st.error(f"PDF 문서 생성 오류: {e}")
            return None
    
    def create_hwp_report(self, content: str, title: str = "서울 상권 분석 보고서") -> Optional[io.BytesIO]:
        """Create an HWP document from the report content."""
        if not hwp:
            st.error("HWP 라이브러리가 설치되지 않았습니다.")
            return None
            
        try:
            # Create temporary HWP file
            with tempfile.NamedTemporaryFile(suffix='.hwp', delete=False) as tmp_file:
                tmp_path = tmp_file.name
            
            # Create HWP document
            doc = hwp.HwpDocument()
            doc.open(tmp_path)
            
            # Add title
            doc.insert_text(title)
            doc.insert_paragraph()
            doc.insert_text(f"생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}")
            doc.insert_paragraph()
            
            # Add content
            self._parse_content_to_hwp(doc, content)
            
            # Save and read
            doc.save()
            doc.close()
            
            # Read the file
            with open(tmp_path, 'rb') as f:
                content_bytes = f.read()
            
            # Clean up
            os.unlink(tmp_path)
            
            bio = io.BytesIO(content_bytes)
            bio.seek(0)
            return bio
            
        except Exception as e:
            st.error(f"HWP 문서 생성 오류: {e}")
            return None
    
    def _parse_content_to_word(self, document, content: str):
        """Parse content and add to Word document."""
        for line in content.split('\n'):
            line = line.strip()
            if line.startswith('# '):
                document.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                document.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                document.add_heading(line[4:], level=3)
            elif line.startswith('- '):
                document.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. '):
                document.add_paragraph(line[3:], style='List Number')
            elif line:
                document.add_paragraph(line)
            else:
                document.add_paragraph()
    
    def _parse_content_to_pdf(self, pdf, content: str):
        """Parse content and add to PDF document."""
        for line in content.split('\n'):
            line = line.strip()
            if line.startswith('# '):
                pdf.set_font('Korean' if 'Korean' in pdf.fonts else 'Arial', 'B', 14)
                pdf.cell(0, 10, line[2:], ln=True)
            elif line.startswith('## '):
                pdf.set_font('Korean' if 'Korean' in pdf.fonts else 'Arial', 'B', 12)
                pdf.cell(0, 8, line[3:], ln=True)
            elif line.startswith('### '):
                pdf.set_font('Korean' if 'Korean' in pdf.fonts else 'Arial', 'B', 11)
                pdf.cell(0, 6, line[4:], ln=True)
            elif line.startswith('- '):
                pdf.set_font('Korean' if 'Korean' in pdf.fonts else 'Arial', '', 10)
                pdf.cell(10, 6, '•', ln=False)
                pdf.cell(0, 6, line[2:], ln=True)
            elif line.startswith('1. '):
                pdf.set_font('Korean' if 'Korean' in pdf.fonts else 'Arial', '', 10)
                pdf.cell(10, 6, '1.', ln=False)
                pdf.cell(0, 6, line[3:], ln=True)
            elif line:
                pdf.set_font('Korean' if 'Korean' in pdf.fonts else 'Arial', '', 10)
                # Handle long lines
                if len(line) > 80:
                    pdf.multi_cell(0, 6, line)
                else:
                    pdf.cell(0, 6, line, ln=True)
            else:
                pdf.ln(3)
    
    def _parse_content_to_hwp(self, doc, content: str):
        """Parse content and add to HWP document."""
        for line in content.split('\n'):
            line = line.strip()
            if line.startswith('# '):
                # Set title style
                doc.insert_text(line[2:])
                doc.insert_paragraph()
            elif line.startswith('## '):
                # Set heading style
                doc.insert_text(line[3:])
                doc.insert_paragraph()
            elif line.startswith('### '):
                # Set subheading style
                doc.insert_text(line[4:])
                doc.insert_paragraph()
            elif line.startswith('- '):
                # Add bullet point
                doc.insert_text(f"• {line[2:]}")
                doc.insert_paragraph()
            elif line.startswith('1. '):
                # Add numbered list
                doc.insert_text(f"1. {line[3:]}")
                doc.insert_paragraph()
            elif line:
                doc.insert_text(line)
                doc.insert_paragraph()
            else:
                doc.insert_paragraph()
    
    def _get_korean_font_path(self) -> Optional[str]:
        """Get Korean font path."""
        possible_paths = [
            'NanumGothic.ttf',
            'malgun.ttf',
            'gulim.ttc',
            'C:/Windows/Fonts/malgun.ttf',
            'C:/Windows/Fonts/gulim.ttc',
            'C:/Windows/Fonts/NanumGothic.ttf'
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        return None
    
    def get_file_info(self, format_type: str) -> Dict[str, str]:
        """Get file information for download."""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"서울상권분석보고서_{timestamp}"
        
        file_info = {
            'pdf': {
                'filename': f"{base_filename}.pdf",
                'mime_type': 'application/pdf',
                'label': '📄 PDF 다운로드',
                'help': 'PDF 형식으로 보고서를 다운로드합니다.'
            },
            'docx': {
                'filename': f"{base_filename}.docx",
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'label': '📄 Word 다운로드',
                'help': 'Microsoft Word 형식으로 보고서를 다운로드합니다.'
            },
            'hwp': {
                'filename': f"{base_filename}.hwp",
                'mime_type': 'application/octet-stream',
                'label': '📄 HWP 다운로드',
                'help': '한글(HWP) 형식으로 보고서를 다운로드합니다.'
            }
        }
        
        return file_info.get(format_type, {})

def render_download_buttons(content: str, title: str = "서울 상권 분석 보고서"):
    """Render download buttons for the report."""
    if not content:
        st.warning("다운로드할 보고서 내용이 없습니다.")
        return
    
    downloader = ReportDownloader()
    
    st.markdown("---")
    st.markdown("### 📥 보고서 다운로드")
    st.markdown("원하는 형식으로 보고서를 다운로드하세요.")
    
    # Create columns for download buttons
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # PDF Download
        pdf_info = downloader.get_file_info('pdf')
        try:
            pdf_bio = downloader.create_pdf_report(content, title)
            if pdf_bio:
                st.download_button(
                    label=pdf_info['label'],
                    data=pdf_bio,
                    file_name=pdf_info['filename'],
                    mime=pdf_info['mime_type'],
                    use_container_width=True,
                    help=pdf_info['help']
                )
        except Exception as e:
            st.error(f"PDF 생성 오류: {e}")
    
    with col2:
        # Word Download
        word_info = downloader.get_file_info('docx')
        try:
            word_bio = downloader.create_word_report(content, title)
            if word_bio:
                st.download_button(
                    label=word_info['label'],
                    data=word_bio,
                    file_name=word_info['filename'],
                    mime=word_info['mime_type'],
                    use_container_width=True,
                    help=word_info['help']
                )
        except Exception as e:
            st.error(f"Word 생성 오류: {e}")
    
    with col3:
        # HWP Download
        hwp_info = downloader.get_file_info('hwp')
        if 'hwp' in downloader.supported_formats:
            try:
                hwp_bio = downloader.create_hwp_report(content, title)
                if hwp_bio:
                    st.download_button(
                        label=hwp_info['label'],
                        data=hwp_bio,
                        file_name=hwp_info['filename'],
                        mime=hwp_info['mime_type'],
                        use_container_width=True,
                        help=hwp_info['help']
                    )
            except Exception as e:
                st.error(f"HWP 생성 오류: {e}")
        else:
            st.info("📄 HWP 지원 예정")
            st.caption("한글(HWP) 형식은 추후 지원될 예정입니다.")

# Global instance
report_downloader = ReportDownloader()
