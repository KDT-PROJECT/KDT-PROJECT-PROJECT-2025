"""
보고서 내보내기 유틸리티
PDF, Word, Excel 형식으로 보고서를 내보내는 기능을 제공합니다.
"""

import io
import logging
from datetime import datetime
from typing import Optional, Dict, Any
import streamlit as st

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

try:
    import pandas as pd
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils.dataframe import dataframe_to_rows
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

logger = logging.getLogger(__name__)

class ReportExporter:
    """보고서 내보내기 클래스"""
    
    def __init__(self):
        self.pdf_available = PDF_AVAILABLE
        self.word_available = WORD_AVAILABLE
        self.excel_available = EXCEL_AVAILABLE
        
    def create_download_buttons(self, markdown_content: str, title: str = "보고서"):
        """다운로드 버튼들을 생성합니다."""
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # PDF 다운로드 버튼
        if self.pdf_available:
            try:
                pdf_data = self._create_pdf(markdown_content, title)
                st.download_button(
                    label="📄 PDF 다운로드",
                    data=pdf_data,
                    file_name=f"{title}_{timestamp}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"PDF 생성 오류: {str(e)}")
        else:
            st.warning("⚠️ PDF 다운로드 불가 (ReportLab 미설치)")
        
        # Word 다운로드 버튼
        if self.word_available:
            try:
                word_data = self._create_word(markdown_content, title)
                st.download_button(
                    label="📝 Word 다운로드",
                    data=word_data,
                    file_name=f"{title}_{timestamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Word 생성 오류: {str(e)}")
        else:
            st.warning("⚠️ Word 다운로드 불가 (python-docx 미설치)")
        
        # Markdown 다운로드 버튼 (항상 사용 가능)
        st.download_button(
            label="📋 Markdown 다운로드",
            data=markdown_content,
            file_name=f"{title}_{timestamp}.md",
            mime="text/markdown",
            use_container_width=True
        )
    
    def _create_pdf(self, markdown_content: str, title: str) -> bytes:
        """PDF 보고서를 생성합니다."""
        buffer = io.BytesIO()
        
        # PDF 문서 생성
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, 
                               topMargin=72, bottomMargin=18)
        
        # 스타일 정의
        styles = getSampleStyleSheet()
        
        # 제목 스타일
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        # 부제목 스타일
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue
        )
        
        # 본문 스타일
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            alignment=TA_LEFT
        )
        
        # 스토리 구성
        story = []
        
        # 제목 추가
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # 생성 날짜
        story.append(Paragraph(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}", body_style))
        story.append(Spacer(1, 20))
        
        # 마크다운 내용을 파싱하여 PDF로 변환
        lines = markdown_content.split('\n')
        current_section = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
                
            if line.startswith('# '):
                # 메인 제목
                story.append(Paragraph(line[2:], title_style))
                story.append(Spacer(1, 12))
            elif line.startswith('## '):
                # 부제목
                story.append(Paragraph(line[3:], subtitle_style))
                story.append(Spacer(1, 8))
            elif line.startswith('### '):
                # 소제목
                story.append(Paragraph(line[4:], subtitle_style))
                story.append(Spacer(1, 6))
            elif line.startswith('- ') or line.startswith('* '):
                # 리스트 항목
                story.append(Paragraph(f"• {line[2:]}", body_style))
            elif line.startswith('**') and line.endswith('**'):
                # 굵은 글씨
                story.append(Paragraph(line, body_style))
            else:
                # 일반 텍스트
                if line:
                    story.append(Paragraph(line, body_style))
        
        # PDF 생성
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _create_word(self, markdown_content: str, title: str) -> bytes:
        """Word 보고서를 생성합니다."""
        buffer = io.BytesIO()
        
        # Word 문서 생성
        doc = Document()
        
        # 제목 추가
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 생성 날짜
        date_para = doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 빈 줄 추가
        doc.add_paragraph()
        
        # 마크다운 내용을 파싱하여 Word로 변환
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
                
            if line.startswith('# '):
                # 메인 제목
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # 부제목
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # 소제목
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                # 리스트 항목
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('**') and line.endswith('**'):
                # 굵은 글씨
                para = doc.add_paragraph()
                run = para.add_run(line[2:-2])
                run.bold = True
            else:
                # 일반 텍스트
                if line:
                    doc.add_paragraph(line)
        
        # Word 문서 저장
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def get_installation_guide(self) -> str:
        """설치 가이드를 반환합니다."""
        guide = """
# 보고서 내보내기 라이브러리 설치 가이드

## PDF 다운로드 (ReportLab)
pip install reportlab

## Word 다운로드 (python-docx)
pip install python-docx

## Excel 다운로드 (openpyxl)
pip install openpyxl

## 전체 설치
pip install reportlab python-docx openpyxl
        """
        return guide.strip()

# 전역 인스턴스
_report_exporter = None

def get_report_exporter() -> ReportExporter:
    """보고서 내보내기 인스턴스를 반환합니다."""
    global _report_exporter
    if _report_exporter is None:
        _report_exporter = ReportExporter()
    return _report_exporter

if __name__ == "__main__":
    # 테스트
    exporter = get_report_exporter()
    print(f"PDF 사용 가능: {exporter.pdf_available}")
    print(f"Word 사용 가능: {exporter.word_available}")
    print(f"Excel 사용 가능: {exporter.excel_available}")