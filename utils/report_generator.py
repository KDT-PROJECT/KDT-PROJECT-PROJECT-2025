"""
보고서 생성 및 다운로드 모듈
PRD TASK3: HWP, DOCX, PDF 형식 보고서 생성
"""

import logging
import os
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import pandas as pd
from io import BytesIO
import base64

# Document generation libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT

logger = logging.getLogger(__name__)

class McKinseyReportGenerator:
    """맥킨지 스타일 보고서 생성기"""
    
    def __init__(self):
        """보고서 생성기 초기화"""
        self.output_dir = Path("reports")
        self.output_dir.mkdir(exist_ok=True)
    
    def generate_markdown_report(self, df: pd.DataFrame, data: Dict[str, Any], web_results: List[Dict] = None) -> str:
        """마크다운 보고서 생성"""
        try:
            report = f"""# 🏙️ 서울 상권분석 보고서
## Seoul Commercial Analysis Report

**생성일시**: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}  
**분석 기간**: 2024년 1월 - 12월  
**보고서 버전**: v1.0  

---

## 📊 경영진 요약 (Executive Summary)

### 핵심 성과 지표
- **총 매출**: ₩{data.get('total_sales', 0):,.0f}
- **성장률**: {data.get('growth_rate', 0):.1f}%
- **분석 지역 수**: {data.get('region_count', 0)}개
- **분석 업종 수**: {data.get('industry_count', 0)}개

### 주요 인사이트
1. **지역별 성과**: {data.get('top_region', '강남구')}가 최고 성과를 보임
2. **업종별 트렌드**: {data.get('top_industry', '소매업')}이 주도적 성장
3. **성장 동력**: 디지털 전환과 고객 경험 개선이 핵심 성장 요인

---

## 📈 상세 분석 (Detailed Analysis)

### 데이터 개요
- **분석 기간**: 2024년 1월 - 12월
- **총 데이터 포인트**: {len(df):,}개
- **분석 대상**: 서울시 주요 상권 및 업종

---

## 🎯 전략적 권고사항 (Strategic Recommendations)

### 1. 즉시 실행 가능한 개선사항
- **고성과 지역 모델 확산**: {data.get('top_region', '강남구')}의 성공 요인을 다른 지역에 적용
- **디지털 전환 가속화**: 온라인-오프라인 통합 서비스 강화

### 2. 중기 전략 (6-12개월)
- **신규 시장 진출**: 성장 잠재력이 높은 지역 발굴 및 진출
- **업종 다각화**: {data.get('top_industry', '소매업')} 외 신규 업종 진출

---

*본 보고서는 서울 상권분석 LLM 시스템에서 자동 생성되었습니다.*
            """
            
            logger.info("마크다운 보고서 생성 완료")
            return report

        except Exception as e:
            logger.error(f"마크다운 보고서 생성 실패: {str(e)}")
            return "# 보고서 생성 중 오류가 발생했습니다."
    
    def generate_docx_report(self, df: pd.DataFrame, data: Dict[str, Any], web_results: List[Dict] = None) -> bytes:
        """DOCX 보고서 생성"""
        try:
            doc = Document()
            
            # 제목
            title = doc.add_heading('서울 상권분석 보고서', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 메타 정보
            meta_info = doc.add_paragraph()
            meta_info.add_run(f"생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}\n")
            meta_info.add_run(f"분석 기간: 2024년 1월 - 12월\n")
            meta_info.add_run(f"보고서 버전: v1.0")
            
            # 경영진 요약
            doc.add_heading('📊 경영진 요약 (Executive Summary)', level=1)
            doc.add_paragraph(f"총 매출: ₩{data.get('total_sales', 0):,.0f}")
            doc.add_paragraph(f"성장률: {data.get('growth_rate', 0):.1f}%")
            doc.add_paragraph(f"분석 지역 수: {data.get('region_count', 0)}개")
            doc.add_paragraph(f"분석 업종 수: {data.get('industry_count', 0)}개")
            
            # 상세 분석
            doc.add_heading('📈 상세 분석 (Detailed Analysis)', level=1)
            doc.add_paragraph(f"총 데이터 포인트: {len(df):,}개")
            doc.add_paragraph("분석 대상: 서울시 주요 상권 및 업종")
            
            # 권고사항
            doc.add_heading('🎯 전략적 권고사항 (Strategic Recommendations)', level=1)
            doc.add_paragraph("1. 고성과 지역 모델 확산")
            doc.add_paragraph("2. 디지털 전환 가속화")
            doc.add_paragraph("3. 신규 시장 진출")
            
            # 바이트 스트림으로 저장
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            logger.info("DOCX 보고서 생성 완료")
            return doc_bytes.getvalue()

        except Exception as e:
            logger.error(f"DOCX 보고서 생성 실패: {str(e)}")
            return b""
    
    def generate_pdf_report(self, df: pd.DataFrame, data: Dict[str, Any], web_results: List[Dict] = None) -> bytes:
        """PDF 보고서 생성"""
        try:
            pdf_bytes = BytesIO()
            doc = SimpleDocTemplate(pdf_bytes, pagesize=A4)
            
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            story = []
            
            # 제목
            story.append(Paragraph("서울 상권분석 보고서", title_style))
            story.append(Paragraph("Seoul Commercial Analysis Report", title_style))
            story.append(Spacer(1, 20))
            
            # 메타 정보
            meta_text = f"""
            생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}<br/>
            분석 기간: 2024년 1월 - 12월<br/>
            보고서 버전: v1.0
            """
            story.append(Paragraph(meta_text, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # 경영진 요약
            story.append(Paragraph("📊 경영진 요약 (Executive Summary)", styles['Heading2']))
            story.append(Paragraph(f"총 매출: ₩{data.get('total_sales', 0):,.0f}", styles['Normal']))
            story.append(Paragraph(f"성장률: {data.get('growth_rate', 0):.1f}%", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # PDF 생성
            doc.build(story)
            pdf_bytes.seek(0)
            
            logger.info("PDF 보고서 생성 완료")
            return pdf_bytes.getvalue()

        except Exception as e:
            logger.error(f"PDF 보고서 생성 실패: {str(e)}")
            return b""

def generate_report_downloads(df: pd.DataFrame, data: Dict[str, Any], web_results: List[Dict] = None) -> Dict[str, bytes]:
    """모든 형식의 보고서 다운로드 생성"""
    generator = McKinseyReportGenerator()
    
    reports = {
        "markdown": generator.generate_markdown_report(df, data, web_results).encode('utf-8'),
        "docx": generator.generate_docx_report(df, data, web_results),
        "pdf": generator.generate_pdf_report(df, data, web_results)
    }
    
    return reports