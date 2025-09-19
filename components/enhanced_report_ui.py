"""Enhanced Report Generation UI Components"""

import streamlit as st
import pandas as pd
from typing import Dict, Any, List
from datetime import datetime
import io
from docx import Document
from fpdf import FPDF

class EnhancedReportUI:
    """Enhanced UI components for report generation"""

    @staticmethod
    def _create_word_report(content: str) -> io.BytesIO:
        """Create a Word document from the report content."""
        document = Document()
        for line in content.split('\n'):
            if line.startswith('# '):
                document.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                document.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                document.add_heading(line[4:], level=3)
            elif line.strip():
                document.add_paragraph(line)
        bio = io.BytesIO()
        document.save(bio)
        bio.seek(0)
        return bio

    @staticmethod
    def _create_pdf_report(content: str) -> io.BytesIO:
        """Create a PDF document from the report content."""
        pdf = FPDF()
        pdf.add_page()
        
        try:
            pdf.add_font('NanumGothic', '', 'NanumGothic.ttf', uni=True)
            pdf.set_font('NanumGothic', '', 12)
        except RuntimeError:
            pdf.set_font('Arial', '', 12)
            st.warning("한글 폰트(NanumGothic.ttf)를 찾을 수 없어 PDF 내보내기 시 한글이 깨질 수 있습니다.")

        for line in content.split('\n'):
            line_text = line.strip()
            if line_text.startswith('# '):
                pdf.set_font('NanumGothic', 'B', 16)
                pdf.cell(0, 10, line_text[2:], ln=True)
            elif line_text.startswith('## '):
                pdf.set_font('NanumGothic', 'B', 14)
                pdf.cell(0, 10, line_text[3:], ln=True)
            elif line_text.startswith('### '):
                pdf.set_font('NanumGothic', 'B', 12)
                pdf.cell(0, 10, line_text[4:], ln=True)
            elif line_text:
                pdf.set_font('NanumGothic', '', 12)
                pdf.multi_cell(0, 10, line_text)
            else:
                pdf.ln(5)
        
        pdf_output = pdf.output(dest='S').encode('latin-1')
        bio = io.BytesIO(pdf_output)
        bio.seek(0)
        return bio

    @staticmethod
    def render_report_hero():
        """Render hero section for report generation"""
        st.markdown("""
        <div style="text-align: center; padding: 2rem 0;">
            <h2>📋 AI 기반 전문 보고서 생성</h2>
            <p>데이터 분석과 시각화를 통해 맥킨지 스타일의 전문 보고서를 생성합니다.</p>
        </div>
        """, unsafe_allow_html=True)

    @staticmethod
    def render_data_status():
        """Render data availability status"""
        sql_available = st.session_state.get('last_sql_df') is not None
        rag_available = st.session_state.get('last_web_results') is not None
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("SQL 데이터", "✅ 사용 가능" if sql_available else "❌ 없음")
        with col2:
            st.metric("웹 검색 데이터", "✅ 사용 가능" if rag_available else "❌ 없음")
            
        return {"sql_available": sql_available, "rag_available": rag_available}

    @staticmethod
    def render_generated_report(report_data: Dict[str, Any], options: Dict[str, Any]):
        """Render the generated report, including charts."""
        if not report_data.get("status") == "success":
            st.error(f"보고서 생성 실패: {report_data.get('message', '알 수 없는 오류')}")
            return

        st.success("🎉 보고서 생성이 완료되었습니다!")
        
        report_content = report_data.get("content", "보고서 내용을 불러올 수 없습니다.")
        charts = report_data.get("charts", [])

        # Display Charts first for visual impact
        if charts:
            st.markdown("### 📊 데이터 시각화")
            for chart in charts:
                st.plotly_chart(chart, use_container_width=True)
        
        # Display McKinsey-style report content
        with st.container():
            st.markdown("### 📄 맥킨지 스타일 보고서")
            st.markdown(report_content, unsafe_allow_html=True)

        # Render download buttons
        EnhancedReportUI._render_download_options(report_content)

        # Metadata and additional info
        with st.expander("📋 보고서 상세 정보 보기"):
            col1, col2 = st.columns(2)
            with col1:
                if options.get("include_metadata") and report_data.get("metadata"):
                    st.subheader("메타데이터")
                    st.json(report_data["metadata"])
            with col2:
                if report_data.get("kpis"):
                    st.subheader("주요 성과 지표 (KPI)")
                    st.json(report_data["kpis"])

    @staticmethod
    def _render_download_options(report_content: str):
        """Render download buttons for the report."""
        st.markdown("---")
        st.markdown("### 📥 보고서 다운로드")

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"mckinsey_style_report_{timestamp}"

        col1, col2, col3 = st.columns(3)

        with col1:
            try:
                word_bio = EnhancedReportUI._create_word_report(report_content)
                st.download_button(
                    label="📄 Word (.docx)",
                    data=word_bio,
                    file_name=f"{base_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Word 변환 오류: {e}")

        with col2:
            try:
                pdf_bio = EnhancedReportUI._create_pdf_report(report_content)
                st.download_button(
                    label="📄 PDF (.pdf)",
                    data=pdf_bio,
                    file_name=f"{base_filename}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"PDF 변환 오류: {e}")

        with col3:
            if st.button("📄 HWP (.hwp)", use_container_width=True, help="한글(HWP) 파일 다운로드는 현재 지원되지 않습니다."):
                st.info("한글(HWP) 파일 형식은 현재 지원되지 않으며, 추후 지원될 예정입니다.")
